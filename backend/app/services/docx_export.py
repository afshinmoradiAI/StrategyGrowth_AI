"""DOCX export for a completed plan."""
from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor


def _h1(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)


def _h2(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def render_plan_docx(plan: dict[str, Any]) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    brief = plan.get("brief") or {}
    research = plan.get("research") or {}
    strategy = plan.get("strategy") or {}
    roadmap = plan.get("roadmap") or {}
    risks = (plan.get("risks") or {}).get("risks", []) if plan.get("risks") else []

    title = brief.get("project_name") or "Strategic Plan"
    p = doc.add_heading(title, level=0)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    if brief.get("domain"):
        doc.add_paragraph(brief["domain"]).italic = True
    if brief.get("summary"):
        doc.add_paragraph(brief["summary"])

    if brief.get("goals"):
        _h2(doc, "Goals")
        _bullets(doc, brief["goals"])
    if brief.get("target_audience"):
        _h2(doc, "Target Audience")
        _bullets(doc, brief["target_audience"])
    if brief.get("constraints"):
        _h2(doc, "Constraints")
        _bullets(doc, brief["constraints"])
    if brief.get("success_criteria"):
        _h2(doc, "Success Criteria")
        _bullets(doc, brief["success_criteria"])

    if research:
        doc.add_page_break()
        _h1(doc, "Market Research")
        if research.get("market_overview"):
            doc.add_paragraph(research["market_overview"])
        if research.get("trends"):
            _h2(doc, "Trends")
            _bullets(doc, research["trends"])
        if research.get("competitors"):
            _h2(doc, "Competitors")
            for c in research["competitors"]:
                p = doc.add_paragraph()
                p.add_run(c.get("name", "")).bold = True
                if c.get("description"):
                    p.add_run(f" — {c['description']}")

    if strategy:
        doc.add_page_break()
        _h1(doc, "Go-To-Market Strategy")
        if strategy.get("positioning"):
            p = doc.add_paragraph(strategy["positioning"])
            p.italic = True
        if strategy.get("value_propositions"):
            _h2(doc, "Value Propositions")
            _bullets(doc, strategy["value_propositions"])
        if strategy.get("differentiators"):
            _h2(doc, "Differentiators")
            _bullets(doc, strategy["differentiators"])
        if strategy.get("objectives"):
            _h2(doc, "Objectives")
            for o in strategy["objectives"]:
                p = doc.add_paragraph()
                p.add_run(o.get("objective", "")).bold = True
                if o.get("rationale"):
                    doc.add_paragraph(o["rationale"])

    if roadmap and roadmap.get("phases"):
        doc.add_page_break()
        _h1(doc, "Roadmap")
        for i, phase in enumerate(roadmap["phases"], 1):
            _h2(doc, f"Phase {i}: {phase.get('name', '')}")
            if phase.get("duration"):
                doc.add_paragraph(f"Duration: {phase['duration']}").italic = True
            if phase.get("objective"):
                doc.add_paragraph(phase["objective"])
            for m in phase.get("milestones", []):
                p = doc.add_paragraph()
                p.add_run(f"• {m.get('title', '')}").bold = True
                if m.get("target"):
                    p.add_run(f" ({m['target']})")
                for t in m.get("tasks", []):
                    bullet = f"   – {t.get('title', '')}"
                    if t.get("estimate"):
                        bullet += f" [{t['estimate']}]"
                    doc.add_paragraph(bullet)

    if risks:
        doc.add_page_break()
        _h1(doc, "Risk Register")
        for r in risks:
            p = doc.add_paragraph()
            p.add_run(f"{r.get('id', '')} ").bold = True
            p.add_run(r.get("description", ""))
            doc.add_paragraph(
                f"Likelihood: {r.get('likelihood', '-')} · "
                f"Impact: {r.get('impact', '-')} · "
                f"Category: {r.get('category', '-')}"
            )
            if r.get("mitigation"):
                doc.add_paragraph(f"Mitigation: {r['mitigation']}")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
